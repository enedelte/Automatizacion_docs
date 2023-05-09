# Inportar libreria DOCX
from docx import Document
# Esta llave es el documento nuevo creado
document = Document()
# Crea titulo 1
document.add_heading("Hola mundo")
# (, Level=2; Level=3) es el separador en titulos heading
document.add_heading("este es mi primer mensaje", level=2)
document.add_heading("para todos", level=3)
paragraph = document.add_paragraph("este es otro mensaje que envio")
paragraph.add_run(" y esta letra está en negrilla.").bold = True
paragraph.add_run(" y esta letra está en italica.").italic = True
# document.add_page_break divide páginas
document.add_page_break()

# Se guarda documento con el nombre ("")
document.save("prueba1.docx") 


#print(Document)
